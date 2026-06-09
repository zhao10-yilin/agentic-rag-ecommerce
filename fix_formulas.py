"""
Fix script for formula numbering and missing table title
Formulas are OMML math objects - need to add equation numbers as runs
"""
import sys
import io
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

def add_equation_number_to_omath_para(para, eq_num):
    """Add a right-aligned equation number to an OMML math paragraph"""
    # Add a run with tab + equation number
    run = para.add_run(f'\t({eq_num})')
    run.font.size = Pt(10.5)  # Match body text size

    # Add right-aligned tab stop to paragraph properties
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(para._element, qn('w:pPr'))

    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = etree.SubElement(pPr, qn('w:tabs'))

    # Add right-aligned tab stop at page right margin
    # Standard A4 page: ~15.6cm text width (210mm - 25.4mm*2 margins)
    tab_elem = etree.SubElement(tabs, qn('w:tab'))
    tab_elem.set(qn('w:val'), 'right')
    tab_elem.set(qn('w:pos'), '9072')  # ~16cm in twips (9072 = 16cm)

def fix_formula_numbering(doc_path, output_path):
    doc = Document(doc_path)

    # Map paragraph indices to formula numbers
    formula_map = {
        84: '2-1',   # Fick's law
        87: '2-2',   # Page model
        97: '2-3',   # Carnot COP
        102: '2-4',  # COP definition
        105: '2-5',  # Energy balance
        108: '2-6',  # Dehumidification
        193: '3-1',  # Total heat load
        200: '3-2',  # COP and SMER
        214: '4-1',  # Theoretical compression power
        216: '4-2',  # Actual compression power
        226: '4-3',  # Heat transfer equation
        228: '4-4',  # LMTD
        238: '4-5',  # Fan shaft power
        240: '4-6',  # System resistance
    }

    count = 0
    for idx, eq_num in sorted(formula_map.items()):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            # Check if this is an OMML paragraph
            xml = para._element.xml
            if 'oMathPara' in xml or 'oMath' in xml:
                # Check if we already added the number
                already_added = False
                for run in para.runs:
                    if eq_num in run.text:
                        already_added = True
                        break

                if not already_added:
                    add_equation_number_to_omath_para(para, eq_num)
                    count += 1
                    print(f"  P{idx}: 添加公式编号 ({eq_num}) [OMML]")
            else:
                # Regular text paragraph - add as text
                if para.runs:
                    # Check if already added
                    already_added = False
                    for run in para.runs:
                        if eq_num in run.text:
                            already_added = True
                            break
                    if not already_added:
                        para.add_run(f'    ({eq_num})')
                        count += 1
                        print(f"  P{idx}: 添加公式编号 ({eq_num}) [text]")

    print(f"  共为 {count} 个公式添加了编号")

    # Also add formula references in the text after formula descriptions
    # These are the "其中" paragraphs that explain formula variables
    ref_additions = {
        85: '式(2-1)中',     # After Fick's law: "其中，MR是含水率..."
        88: '式(2-2)中',     # After Page model
        98: '式(2-3)中',     # After Carnot COP
        103: '式(2-4)中',    # After COP formula
        106: '式(2-5)中',    # After energy balance
        109: '式(2-6)中',    # After dehumidification
        195: '式(3-1)中',    # After heat load
        202: '式(3-2)中',    # After COP/SMER
        215: '式(4-1)中',    # After compression power
        217: '式(4-2)中',    # After actual power
        227: '式(4-3)中',    # After heat transfer
        229: '式(4-4)中',    # After LMTD
        239: '式(4-5)中',    # After fan power
        241: '式(4-6)中',    # After system resistance
    }

    ref_count = 0
    for idx, prefix in sorted(ref_additions.items()):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            if para.runs and not para.runs[0].text.startswith(prefix):
                para.runs[0].text = prefix + para.runs[0].text
                ref_count += 1
                if ref_count <= 3:
                    print(f"  P{idx}: 添加公式引用前缀 '{prefix}'")

    print(f"  共为 {ref_count} 个说明段落添加了公式引用")

    # Fix missing P126 table title
    # Find the table title paragraph for 除湿转轮耦合技术
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if '除湿转轮耦合技术性能参数对比' in text and '表2-1' not in text and '表2-2' not in text:
            # Check if still has old numbering
            if '表2.1' in text:
                new_text = text.replace('表2.1', '表2-2')
                if p.runs:
                    p.runs[0].text = new_text
                print(f"  P{i}: 修正表格标题 → '表2-2'")
                break

    # Save
    doc.save(output_path)
    print(f"\n修正完成，保存至: {output_path}")

if __name__ == '__main__':
    input_file = '广东工业大学张智朝毕业论文平菇热泵烘干房设计 - 修订版.docx'
    output_file = '广东工业大学张智朝毕业论文平菇热泵烘干房设计 - 修订版.docx'
    fix_formula_numbering(input_file, output_file)

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from lxml import etree
from docx.shared import Pt

doc = Document('revised_v5.docx')
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Get body element directly
body = doc.element.body

# ============================================================
# Step 1: Find and remove old TOC area
# ============================================================
toc_para = None
toc_idx_in_body = None
for i, child in enumerate(body):
    if child.tag == f'{{{ns_w}}}p':
        pPr = child.find(f'{{{ns_w}}}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{{{ns_w}}}pStyle')
            if pStyle is not None and pStyle.get(f'{{{ns_w}}}val') == '13':  # toc 4 style id
                toc_para = child
                toc_idx_in_body = i
                break

if toc_para is None:
    print('TOC paragraph not found in body!')
    sys.exit(1)

print(f'TOC paragraph found at body index {toc_idx_in_body}')

# Find range of empty paragraphs before TOC (TOC area start)
toc_start = toc_idx_in_body
for i in range(toc_idx_in_body - 1, max(0, toc_idx_in_body - 15), -1):
    child = body[i]
    if child.tag == f'{{{ns_w}}}p':
        text_elems = child.findall(f'.//{{{ns_w}}}t')
        texts = ''.join(t.text or '' for t in text_elems).strip()
        if texts == '':
            toc_start = i
        else:
            break

print(f'Removing body children [{toc_start}..{toc_idx_in_body}]')

# Remove from end to start (to preserve indices)
for i in range(toc_idx_in_body, toc_start - 1, -1):
    body.remove(body[i])

print(f'Removed {toc_idx_in_body - toc_start + 1} elements')

# ============================================================
# Step 2: Collect headings from document (but at XML level)
# ============================================================
headings = []
for child in body:
    if child.tag == f'{{{ns_w}}}p':
        pPr = child.find(f'{{{ns_w}}}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{{{ns_w}}}pStyle')
            if pStyle is not None:
                sid = pStyle.get(f'{{{ns_w}}}val')
                if sid in ('2', '3', '4'):  # Heading 1=2, Heading 2=3, Heading 3=4
                    text_elems = child.findall(f'.//{{{ns_w}}}t')
                    text = ''.join(t.text or '' for t in text_elems).strip()
                    if text:
                        # Skip 摘要 and Abstract (they go before TOC)
                        if text in ('摘要', 'Abstract'):
                            continue
                        level = {'2': 'H1', '3': 'H2', '4': 'H3'}.get(sid, 'H?')
                        headings.append((level, text))

print(f'Collected {len(headings)} headings from body XML')

# ============================================================
# Step 3: Helper - CJK detection
# ============================================================
def is_cjk(ch):
    cp = ord(ch)
    return (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
            0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF or
            0x2F00 <= cp <= 0x2FDF or 0xFE30 <= cp <= 0xFE4F)

def split_text(text):
    if not text:
        return []
    result = []
    buf = text[0]
    buf_cjk = is_cjk(text[0])
    for ch in text[1:]:
        ch_cjk = is_cjk(ch)
        if ch_cjk == buf_cjk:
            buf += ch
        else:
            result.append((buf, buf_cjk))
            buf = ch
            buf_cjk = ch_cjk
    result.append((buf, buf_cjk))
    return result

# ============================================================
# Step 4: Build TOC entries
# ============================================================
FONT_MAP = {'H1': '黑体', 'H2': '宋体', 'H3': '宋体'}

def make_toc_para(cn_font, text_parts, is_bold=False, is_title=False):
    """Create a TOC paragraph element.
    text_parts is a list of (text, is_cjk) tuples.
    """
    para = etree.Element(f'{{{ns_w}}}p')
    pPr = etree.SubElement(para, f'{{{ns_w}}}pPr')

    if is_title:
        jc = etree.SubElement(pPr, f'{{{ns_w}}}jc')
        jc.set(f'{{{ns_w}}}val', 'center')
        sp = etree.SubElement(pPr, f'{{{ns_w}}}spacing')
        sp.set(f'{{{ns_w}}}after', '300')
        sp.set(f'{{{ns_w}}}line', '480')
        sp.set(f'{{{ns_w}}}lineRule', 'auto')
    else:
        sp = etree.SubElement(pPr, f'{{{ns_w}}}spacing')
        sp.set(f'{{{ns_w}}}line', '380')
        sp.set(f'{{{ns_w}}}lineRule', 'auto')
        # Tab for page number
        tabs = etree.SubElement(pPr, f'{{{ns_w}}}tabs')
        tab = etree.SubElement(tabs, f'{{{ns_w}}}tab')
        tab.set(f'{{{ns_w}}}val', 'right')
        tab.set(f'{{{ns_w}}}leader', 'dot')
        tab.set(f'{{{ns_w}}}pos', '9072')

    # Create runs for each text segment
    for seg_text, is_cjk in text_parts:
        r = etree.SubElement(para, f'{{{ns_w}}}r')
        rPr = etree.SubElement(r, f'{{{ns_w}}}rPr')
        rf = etree.SubElement(rPr, f'{{{ns_w}}}rFonts')
        rf.set(f'{{{ns_w}}}eastAsia', cn_font)
        rf.set(f'{{{ns_w}}}ascii', 'Times New Roman')
        rf.set(f'{{{ns_w}}}hAnsi', 'Times New Roman')
        rf.set(f'{{{ns_w}}}cs', 'Times New Roman')
        size_val = '32' if is_title else '24'
        for tag in ('sz', 'szCs'):
            e = etree.SubElement(rPr, f'{{{ns_w}}}{tag}')
            e.set(f'{{{ns_w}}}val', size_val)
        if is_bold:
            etree.SubElement(rPr, f'{{{ns_w}}}b')
        t = etree.SubElement(r, f'{{{ns_w}}}t')
        t.text = seg_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Tab + leader dots (for non-title entries)
    if not is_title:
        tr = etree.SubElement(para, f'{{{ns_w}}}r')
        trPr = etree.SubElement(tr, f'{{{ns_w}}}rPr')
        trf = etree.SubElement(trPr, f'{{{ns_w}}}rFonts')
        trf.set(f'{{{ns_w}}}ascii', 'Times New Roman')
        trf.set(f'{{{ns_w}}}hAnsi', 'Times New Roman')
        trsz = etree.SubElement(trPr, f'{{{ns_w}}}sz')
        trsz.set(f'{{{ns_w}}}val', '24')
        tt = etree.SubElement(tr, f'{{{ns_w}}}t')
        tt.text = '\t'
        tt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return para

# Create TOC elements
toc_elements = []

# Title - use single pre-split parts that keep spaces with Chinese
# '目  录' should be one run to avoid CJK detection splitting on spaces
title_parts = [('目  录', True)]  # Force as single Chinese segment
toc_elements.append(make_toc_para('黑体', title_parts, is_bold=True, is_title=True))

# Blank line
blank = etree.Element(f'{{{ns_w}}}p')
toc_elements.append(blank)

# Headings
for level, text in headings:
    cn_font = FONT_MAP[level]
    is_bold = (level == 'H1')
    parts = split_text(text)
    toc_elements.append(make_toc_para(cn_font, parts, is_bold=is_bold))

print(f'Created {len(toc_elements)} TOC elements')

# ============================================================
# Step 5: Find first CHAPTER heading (skip 摘要, Abstract)
# TOC goes AFTER abstracts, BEFORE Chapter 1
# ============================================================
insert_before = None
for child in body:
    if child.tag == f'{{{ns_w}}}p':
        pPr = child.find(f'{{{ns_w}}}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{{{ns_w}}}pStyle')
            if pStyle is not None and pStyle.get(f'{{{ns_w}}}val') == '2':  # Heading 1
                text_elems = child.findall(f'.//{{{ns_w}}}t')
                text = ''.join(t.text or '' for t in text_elems).strip()
                # Skip 摘要 and Abstract
                if text in ('摘要', 'Abstract'):
                    continue
                insert_before = child
                break

if insert_before is None:
    print('Cannot find first Heading 1 in body!')
    sys.exit(1)

# Get heading text for confirmation
h1_text_elems = insert_before.findall(f'.//{{{ns_w}}}t')
h1_text = ''.join(t.text or '' for t in h1_text_elems).strip()
print(f'First Heading 1: {h1_text}')

# Insert all TOC elements before the first chapter
# addprevious inserts immediately before target, so first element ends up farthest
for elem in toc_elements:
    insert_before.addprevious(elem)

print(f'Inserted {len(toc_elements)} TOC elements before first chapter')

# ============================================================
# Save and verify
# ============================================================
doc.save('revised_v6.docx')
print('Saved to revised_v6.docx')

# Re-read and verify
doc2 = Document('revised_v6.docx')
ns_w2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

print('\n=== Final TOC Verification (first 12 entries) ===')
count = 0
found_toc = False
for p in doc2.paragraphs:
    text = (p.text or '').strip()

    if '目  录' in text and not found_toc:
        found_toc = True
        print(f'TOC title: {text}')
        for r in p.runs:
            rPr = r._element.find(f'{{{ns_w2}}}rPr')
            ea, asc, sv, bd = '-', '-', '?', ''
            if rPr is not None:
                rf = rPr.find(f'{{{ns_w2}}}rFonts')
                sz = rPr.find(f'{{{ns_w2}}}sz')
                b = rPr.find(f'{{{ns_w2}}}b')
                ea = rf.get(f'{{{ns_w2}}}eastAsia', '-') if rf is not None else '-'
                asc = rf.get(f'{{{ns_w2}}}ascii', '-') if rf is not None else '-'
                sv = sz.get(f'{{{ns_w2}}}val', '?') if sz is not None else '?'
                bd = 'B' if b is not None else ''
            print(f'  Run: cn={ea}, en={asc}, sz={sv} {bd}')
        count += 1
        continue

    # TOC entries have tab characters (for leader dots)
    if found_toc and text and '\t' in text:
        runs_info = []
        for r in p.runs:
            rPr = r._element.find(f'{{{ns_w2}}}rPr')
            ea, asc, sv, bd = '-', '-', '?', ''
            if rPr is not None:
                rf = rPr.find(f'{{{ns_w2}}}rFonts')
                sz = rPr.find(f'{{{ns_w2}}}sz')
                b = rPr.find(f'{{{ns_w2}}}b')
                ea = rf.get(f'{{{ns_w2}}}eastAsia', '-') if rf is not None else '-'
                asc = rf.get(f'{{{ns_w2}}}ascii', '-') if rf is not None else '-'
                sv = sz.get(f'{{{ns_w2}}}val', '?') if sz is not None else '?'
                bd = 'B' if b is not None else ''
            runs_info.append(f'{ea}|{asc}|{sv}{bd}')
        clean_text = text.replace('\t', '[...]')
        print(f'{clean_text[:60]}')
        print(f'  fonts: {runs_info}')
        count += 1
        if count > 12:
            break

    # Stop if we hit Heading 1 (past TOC)
    if found_toc and count > 1 and 'Heading 1' in (p.style.name or ''):
        break

print(f'\nVerified {count} TOC entries')

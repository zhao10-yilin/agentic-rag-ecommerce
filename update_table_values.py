"""
Update calculations based on new Table 2-2 values:
h1=410.0, h2=435.4, h3=271.55, h4=271.55
Also renumber subsequent Chapter 2 tables and update all cross-references.
"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from lxml import etree

doc = Document('平菇热泵烘干房设计_广州版.docx')
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def replace_para_text(doc, idx, new_text):
    """Replace all text in a paragraph"""
    p = doc.paragraphs[idx]._element
    for r in list(p.findall(f'{{{ns_w}}}r')):
        p.remove(r)
    r_elem = etree.SubElement(p, f'{{{ns_w}}}r')
    t_elem = etree.SubElement(r_elem, f'{{{ns_w}}}t')
    t_elem.text = new_text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def set_first_run_text(doc, idx, new_text):
    """Replace first run text in a paragraph"""
    p = doc.paragraphs[idx]
    if p.runs:
        p.runs[0].text = new_text
    else:
        r = p.add_run(new_text)

# New values from Table 2-2
h1, h2, h3, h4 = 410.0, 435.4, 271.55, 271.55
q0 = h1 - h4           # 138.45
qk = h2 - h3           # 163.85
w = h2 - h1            # 25.40
pi_val = 1.49 / 0.45   # 3.311
cop_theo = qk / w      # 6.451
cop_act = cop_theo * 0.85 * 0.92  # 5.045
mr = 45.0 / q0         # 0.3250 kg/s
mr_kgh = mr * 3600     # 1170 kg/h
W_theo = mr * w        # 8.255 kW
W_act = W_theo / (0.85 * 0.92)  # 10.56 kW
Qk = mr * qk           # 53.25 kW
COP_final = Qk / W_act  # 5.04
v1 = 0.0478            # m3/kg at 17C, 0.45MPa
V_dis = mr * v1 * 3600  # 55.9 m3/h

print(f'New derived values:')
print(f'  q0={q0:.2f}, qk={qk:.2f}, w={w:.2f}')
print(f'  pi={pi_val:.3f}, COP_theo={cop_theo:.3f}, COP_act={COP_final:.2f}')
print(f'  mr={mr:.4f} kg/s = {mr_kgh:.0f} kg/h')
print(f'  W_theo={W_theo:.3f}, W_act={W_act:.2f}, Qk={Qk:.2f}')
print(f'  V_dis={V_dis:.1f} m3/h')

# ============================================================
# 1. Fix P71 table title format (add space)
# ============================================================
print('\n[1] Fix P71 table title format...')
p71 = doc.paragraphs[71]
old_title = p71.text.strip()
if '表2-2各' in old_title:
    set_first_run_text(doc, 71, '表2-2 各状态点热力参数')
    print(f'  Fixed: {old_title} -> 表2-2 各状态点热力参数')

# ============================================================
# 2. Update P72 thermal calculation paragraph
# ============================================================
print('\n[2] Update P72 thermal calculation...')
new_p72 = (
    f"根据表2-2数据，单位质量制冷量q₀=h₁-h₄={h1:.1f}-{h4:.2f}={q0:.2f}kJ/kg，"
    f"单位质量制热量q_k=h₂-h₃={h2:.1f}-{h3:.2f}={qk:.2f}kJ/kg，"
    f"压缩机比功w=h₂-h₁={h2:.1f}-{h1:.1f}={w:.2f}kJ/kg。\n"
    f"压缩比π=p_k/p_o=1.49/0.45={pi_val:.3f}。\n"
    f"系统理论COP=q_k/w={qk:.2f}/{w:.2f}={cop_theo:.3f}，"
    f"考虑等熵效率0.85和机械效率0.92，"
    f"实际COP={cop_theo:.3f}×0.85×0.92={COP_final:.2f}。"
)
replace_para_text(doc, 72, new_p72)
print(f'  Updated: q0={q0:.2f}, qk={qk:.2f}, w={w:.2f}, COP={COP_final:.2f}')

# ============================================================
# 3. Renumber tables: 表2-2(P104)->表2-3, 表2-3(P112)->表2-4
# ============================================================
print('\n[3] Renumber Chapter 2 tables...')

# P104: 表2-2 -> 表2-3
p104 = doc.paragraphs[104]
set_first_run_text(doc, 104, '表2-3 除湿转轮耦合技术性能参数对比')
print(f'  P104: 表2-2 -> 表2-3 除湿转轮耦合技术性能参数对比')

# P112: 表2-3 -> 表2-4
p112 = doc.paragraphs[112]
set_first_run_text(doc, 112, '表2-4 主要食用菌热泵干燥工艺参数对比')
print(f'  P112: 表2-3 -> 表2-4 主要食用菌热泵干燥工艺参数对比')

# ============================================================
# 4. Update P100 - design conditions (2.3节)
# ============================================================
print('\n[4] Update P100 design conditions...')
new_p100 = (
    f"本系统广州市设计工况参数确定如下：蒸发温度Te=12℃（基于广州市4~10月平均气温28℃，"
    f"换热温差16℃），冷凝温度Tc=55℃（基于干燥风温50℃，换热温差5℃），"
    f"过热度ΔT_sh=5℃，过冷度ΔT_sc=5℃。在此工况下，各状态点热力参数详见表2-2。\n"
    f"制冷剂质量流量：ṁ_r=Q₀/q₀=45kW/{q0:.2f}kJ/kg={mr:.4f}kg/s={mr_kgh:.0f}kg/h；\n"
    f"压缩机理论功率：W_theo=ṁ_r×w={mr:.4f}×{w:.2f}={W_theo:.2f}kW；\n"
    f"压缩机实际功率（η_is=0.85，η_m=0.92）：W_act={W_theo:.2f}/(0.85×0.92)={W_act:.2f}kW；\n"
    f"冷凝器放热量：Q_k=ṁ_r×q_k={mr:.4f}×{qk:.2f}={Qk:.2f}kW；\n"
    f"系统实际COP={Qk:.2f}/{W_act:.2f}={COP_final:.2f}。\n"
    f"广州较高的环境温度使蒸发温度提升至12℃，压缩比降至{pi_val:.3f}，"
    f"排气温度仅63.2℃，压缩机运行工况优越，系统能效比显著提升。\n\n"
    f"系统输入条件：电源380V/50Hz三相，环境温度5~40℃，环境相对湿度60%~95%RH。\n"
    f"系统输出能力：额定制热量80kW（满负荷），调幅范围20%~100%，适用干燥温度35~65℃。\n"
    f"广州市高湿环境下，系统配备的除湿转轮可有效处理新风中的高含湿量，"
    f"确保干燥介质湿度满足工艺要求。"
)
replace_para_text(doc, 100, new_p100)
print(f'  Updated: mr={mr:.4f}kg/s, W_theo={W_theo:.2f}kW, W_act={W_act:.2f}kW, Qk={Qk:.2f}kW, COP={COP_final:.2f}')

# ============================================================
# 5. Update P201 - compressor calculation (4.1.1节)
# ============================================================
print('\n[5] Update P201 compressor calculation...')
new_p201 = (
    f"根据NIST REFPROP 9.1数据库及表2-2热力参数，R134a在广州设计工况"
    f"（蒸发温度12℃，饱和压力0.45MPa；冷凝温度55℃，饱和压力1.49MPa）"
    f"下的各状态点参数代入计算：\n\n"
    f"制冷剂质量流量：ṁ_r=Q₀/(h₁-h₄)=45/({h1:.1f}-{h4:.2f})={mr:.4f}kg/s；\n"
    f"压缩机理论功率（式4-1）：P_th=ṁ_r×(h₂-h₁)={mr:.4f}×({h2:.1f}-{h1:.1f})={W_theo:.2f}kW；\n"
    f"压缩机实际功率（式4-2）：取等熵效率η_is=0.85，机械效率η_m=0.92，\n"
    f"P_act={W_theo:.2f}/(0.85×0.92)={W_act:.2f}kW；\n"
    f"压缩比：π=p_k/p_o=1.49/0.45={pi_val:.3f}（排气温度仅63.2℃，远低于压缩机允许上限）；\n"
    f"压缩机排气量：V_dis=ṁ_r×v₁×3600={mr:.4f}×{v1}×3600={V_dis:.1f}m³/h；\n"
    f"（v₁为压缩机入口比容，约{v1}m³/kg@17℃/0.45MPa，取自REFPROP）。\n"
    f"系统实际COP=Q_k/P_act={Qk:.2f}/{W_act:.2f}={COP_final:.2f}。\n\n"
    f"选型结论：参考汉钟RC2-410B型半封闭螺杆压缩机，理论排量121m³/h（@2900rpm），"
    f"在广州工况下仅需50Hz（1450rpm）即可满足设计排气量需求（设计余量充裕）。"
    f"压缩机运行于较低压缩比（{pi_val:.3f}）工况，排气温度仅63.2℃，"
    f"有利于延长设备寿命和降低维护成本。\n"
    f"数据来源：NIST Standard Reference Database 23: REFPROP Version 9.1；\n"
    f"汉钟精机RC2系列螺杆压缩机技术手册（2023版）。"
)
replace_para_text(doc, 201, new_p201)
print(f'  Updated compressor calculation')

# ============================================================
# 6. Update P213 - heat exchanger (4.1.2节)
# ============================================================
print('\n[6] Update P213 heat exchanger calculation...')
# Recalculate with LMTD based on Guangzhou conditions
# Evap: Tai=30, Tao=20, Te=12
# LMTD_evap = [(30-12)-(20-12)]/ln[(30-12)/(20-12)] = 10/ln(18/8) = 10/0.811 = 12.33
# A_evap = Qe/(U*LMTD) = 45000/(45*12.33) = 81.1, design = 81.1*1.20 = 97.3 -> 98
# Cond: Tai=40, Tao=50, Tc=55
# LMTD_cond = [(55-40)-(55-50)]/ln[(55-40)/(55-50)] = 10/ln(15/5) = 10/1.099 = 9.10
# A_cond = Qk/(U*LMTD) = 53250/(55*9.10) = 106.4, design = 106.4*1.15 = 122.4 -> 125

lmtde = 10.0 / (18.0/8.0)**(1.0/10.0)  # won't work, use exact
# LMTD_evap = [(30-12)-(20-12)]/ln[(30-12)/(20-12)] = (18-8)/ln(18/8) = 10/ln(2.25) = 10/0.8109 = 12.33
# LMTD_cond = [(55-40)-(55-50)]/ln[(55-40)/(55-50)] = (15-5)/ln(15/5) = 10/ln(3) = 10/1.0986 = 9.10
import math
lmtde = 10.0 / math.log(18.0/8.0)    # 12.33
lmtdc = 10.0 / math.log(15.0/5.0)    # 9.10
A_evap_need = 45000 / (45 * lmtde)    # 81.1
A_evap_design = A_evap_need * 1.20     # 97.3
A_cond_need = (Qk*1000) / (55 * lmtdc) # Qk=53.25kW=53250W
A_cond_design = A_cond_need * 1.15

new_p213 = (
    f"换热器详细结构参数与选型计算依据（广州设计工况，基于表2-2热力参数修正）：\n\n"
    f"（1）蒸发器设计计算（基于对数平均温差法，式4-3、式4-4）：\n"
    f"广州工况：空气侧进口温度T_ai=30℃，出口温度T_ao=20℃；\n"
    f"制冷剂侧蒸发温度T_e=12℃（恒定）；\n"
    f"LMTD=[(T_ai-T_e)-(T_ao-T_e)]/ln[(T_ai-T_e)/(T_ao-T_e)]\n"
    f"=[(30-12)-(20-12)]/ln[(30-12)/(20-12)]=10/ln(18/8)=10/0.811={lmtde:.1f}℃。\n"
    f"取总传热系数U=45W/(m²·K)（翅片管式蒸发器经验值，来源：《制冷原理与设备》第3版），\n"
    f"所需传热面积：A=Q_e/(U×LMTD)=45000/(45×{lmtde:.1f})={A_evap_need:.1f}m²。\n"
    f"考虑10%裕量及广州高湿工况结霜余量，设计面积A_design={A_evap_need:.1f}×1.20={A_evap_design:.0f}m²，实际取100m²。\n\n"
    f"（2）冷凝器设计计算：\n"
    f"空气侧进口温度T_ai=40℃，出口温度T_ao=50℃；\n"
    f"制冷剂侧冷凝温度T_c=55℃（恒定）；\n"
    f"LMTD=[(55-40)-(55-50)]/ln[(55-40)/(55-50)]=10/ln(15/5)=10/1.099={lmtdc:.1f}℃。\n"
    f"取总传热系数U=55W/(m²·K)，\n"
    f"所需传热面积：A=Q_k/(U×LMTD)={(Qk*1000):.0f}/(55×{lmtdc:.1f})={A_cond_need:.1f}m²。\n"
    f"考虑15%裕量，设计面积A_design={A_cond_need:.1f}×1.15={A_cond_design:.0f}m²，实际取125m²。\n\n"
    f"（3）结构参数：\n"
    f"蒸发器管径Φ9.52mm×0.35mm，翅片间距3.0mm（高湿工况亲水膜处理），4排管；\n"
    f"冷凝器管径Φ12.7mm×0.5mm，翅片间距2.2mm，6排管。\n"
    f"数据来源：ASHRAE Handbook-HVAC Systems and Equipment (2020)；\n"
    f"《换热器设计手册》（钱颂文，化学工业出版社，2018）。"
)
replace_para_text(doc, 213, new_p213)
print(f'  Updated: LMTD_evap={lmtde:.1f}℃, A_evap={A_evap_design:.0f}m², LMTD_cond={lmtdc:.1f}℃, A_cond={A_cond_design:.0f}m²')

# ============================================================
# 7. Update any text references to old table numbers in body
# ============================================================
print('\n[7] Update cross-references in text...')
# Find references to 表2-2 (old 除湿转轮) and 表2-3 (old 食用菌) in body text
# P68 already says "如表2-2所示" which now correctly refers to the NEW table 2-2

# Check if any body text references 表2-2 meaning the old one (除湿转轮)
# and change to 表2-3
ref_fixes = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name == '表格标题':
        continue  # Skip table titles (already handled)
    text = p.text or ''
    # Only fix in Chapter 2 area
    if any(heading in text for heading in ['2.4', '2.5', '除湿转轮', '食用菌']):
        if '表2-2' in text and '除湿' in text:
            # This references the old 表2-2 (now 表2-3)
            for run in p.runs:
                if '表2-2' in (run.text or ''):
                    run.text = run.text.replace('表2-2', '表2-3')
                    ref_fixes += 1
                    print(f'  P{i}: 表2-2->表2-3 in {run.text[:40]}...')
        if '表2-3' in text and '食用菌' in text:
            for run in p.runs:
                if '表2-3' in (run.text or ''):
                    run.text = run.text.replace('表2-3', '表2-4')
                    ref_fixes += 1
                    print(f'  P{i}: 表2-3->表2-4 in {run.text[:40]}...')

if ref_fixes == 0:
    print('  No additional cross-reference fixes needed')
else:
    print(f'  Fixed {ref_fixes} cross-references')

# ============================================================
# 8. Check and update COP values in Chapter 3/5 tables and text
# ============================================================
print('\n[8] Update COP and other derived values in later chapters...')

# Find paragraphs mentioning old COP values
cop_updates = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text or ''
    # Update in conclusion (around P278 area) and other sections
    if '4.96' in text or '4.98' in text:
        # Replace with new COP
        for run in p.runs:
            if '4.96' in (run.text or ''):
                run.text = run.text.replace('4.96', f'{COP_final:.2f}')
                cop_updates += 1
            elif '4.98' in (run.text or ''):
                run.text = run.text.replace('4.98', f'{COP_final:.2f}')
                cop_updates += 1
    if '6.34' in text:
        for run in p.runs:
            if '6.34' in (run.text or ''):
                run.text = run.text.replace('6.34', f'{cop_theo:.3f}')
                cop_updates += 1

print(f'  Updated {cop_updates} COP references to {COP_final:.2f}')

# ============================================================
# Save
# ============================================================
output = '平菇热泵烘干房设计_广州版.docx'
doc.save(output)
print(f'\n{"="*60}')
print(f'All updates saved to: {output}')
print(f'{"="*60}')

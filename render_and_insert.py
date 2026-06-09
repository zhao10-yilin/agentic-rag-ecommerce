"""
Render DXF CAD drawings to PNG images and insert them into the thesis
at the correct positions (after corresponding original figures).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np

from ezdxf.addons.drawing import RenderContext, Frontend
from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
from ezdxf import recover
from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os
import glob

cad_dir = 'cad_drawings'
img_dir = 'cad_images'
os.makedirs(img_dir, exist_ok=True)

# ============================================================
# Step 1: Render DXF to PNG
# ============================================================
print('='*60)
print('Step 1: Rendering DXF to PNG...')
print('='*60)

# Try to find a Chinese font for matplotlib
cn_font = None
for fp in font_manager.fontManager.ttflist:
    if 'SimHei' in fp.name or 'SimSun' in fp.name or 'Microsoft YaHei' in fp.name:
        cn_font = fp
        break
if cn_font:
    plt.rcParams['font.sans-serif'] = [cn_font.name, 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    print(f'Using Chinese font: {cn_font.name}')
else:
    print('No Chinese font found, using default')

dxf_files = sorted(glob.glob(os.path.join(cad_dir, 'Fig*.dxf')))

for dxf_path in dxf_files:
    basename = os.path.splitext(os.path.basename(dxf_path))[0]
    png_path = os.path.join(img_dir, f'{basename}.png')

    if os.path.exists(png_path):
        print(f'  {basename}.png already exists, skipping')
        continue

    try:
        doc_dxf, auditor = recover.readfile(dxf_path)
        if auditor.has_errors:
            print(f'  {basename}: DXF has {len(auditor.errors)} errors, attempting recovery...')

        msp = doc_dxf.modelspace()

        fig = plt.figure(figsize=(16, 10), dpi=150)
        ax = fig.add_axes([0.05, 0.05, 0.9, 0.9])

        ctx = RenderContext(doc_dxf)
        # Use default font settings for Chinese compatibility
        out = MatplotlibBackend(ax)
        Frontend(ctx, out).draw_layout(msp, finalize=True)

        ax.autoscale()
        # Remove axes for clean CAD look
        ax.set_axis_off()

        fig.savefig(png_path, dpi=150, bbox_inches='tight',
                    facecolor='white', edgecolor='none', pad_inches=0.1)
        plt.close(fig)
        print(f'  {basename}.png rendered')
    except Exception as e:
        print(f'  {basename}: FAILED - {e}')
        plt.close('all')

print(f'\nPNG images saved to: {img_dir}/')

# ============================================================
# Step 2: First pass - find all insertion positions
# ============================================================
print()
print('='*60)
print('Step 2: Finding insertion positions...')
print('='*60)

doc = Document('平菇热泵烘干房设计_广州版.docx')
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Map: (heading_text, png_name)
fig_mapping = [
    ('3.2.1 热泵主机系统', 'Fig01_HeatPump_System.png', '图3-1 热泵主机系统原理图（CAD）'),
    ('3.2.2 干燥室体设计', 'Fig02_DryingChamber.png', '图3-2 干燥室体结构图（CAD）'),
    ('3.2.3 循环风系统', 'Fig03_AirCirculation.png', '图3-3 循环风系统布置图（CAD）'),
    ('3.2.4 控制系统设计', 'Fig04_PLC_Control.png', '图3-4 PLC控制系统框图（CAD）'),
    ('4.1.1 压缩机选型计算', 'Fig05_R134a_pH_Diagram.png', '图4-1 R134a制冷循环p-h图（CAD）'),
    ('4.1.2 换热器设计选型', 'Fig06_HeatExchanger.png', '图4-2 换热器结构详图（CAD）'),
    ('4.1.3 风机选型与配置', 'Fig07_FanCurve.png', '图4-3 风机性能与管道阻力曲线（CAD）'),
    ('4.2.1 干燥室气流场分析', 'Fig08_AirflowDist.png', '图4-5 干燥室气流速度分布图（CAD）'),
    ('4.2.2 CFD数值模拟建模', 'Fig09_CFD_Mesh.png', '图4-6 CFD模型网格划分图（CAD）'),
    ('4.2.4 结构优化设计', 'Fig10_Optimization.png', '图4-7 干燥室气流组织优化对比图（CAD）'),
    ('4.3.2 安装工艺流程', 'Fig11_Installation.png', '图4-8 安装工艺流程图（CAD）'),
    ('5.1 经济性分析', 'Fig12_Economic.png', '图5-1 热泵烘干房经济性对比分析图（CAD）'),
]

# Find all insertion positions first
insertions = []  # [(img_idx, png_path, caption)]

for heading_text, png_name, caption in fig_mapping:
    png_path = os.path.join(img_dir, png_name)
    if not os.path.exists(png_path):
        print(f'  SKIP {png_name}: PNG not found')
        continue

    # Find heading
    heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or '').strip()
        if 'Heading' in (p.style.name or '') and heading_text in text:
            heading_idx = i
            break

    if heading_idx is None:
        print(f'  SKIP {heading_text}: heading not found')
        continue

    # Find image paragraph in the section
    img_idx = None
    for i in range(heading_idx, min(heading_idx + 30, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        drawings = p._element.findall(f'.//{{{ns_w}}}drawing')
        if drawings and not img_idx:
            img_idx = i
            break

    if img_idx is None:
        # Try wider search
        for i in range(heading_idx, min(heading_idx + 50, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            drawings = p._element.findall(f'.//{{{ns_w}}}drawing')
            if drawings and not img_idx:
                img_idx = i
                break

    if img_idx is None:
        print(f'  SKIP {heading_text}: no image found')
        continue

    # Find insert position: empty para after the image
    insert_idx = img_idx + 1
    for j in range(img_idx + 1, min(img_idx + 5, len(doc.paragraphs))):
        p = doc.paragraphs[j]
        text = (p.text or '').strip()
        drawings = p._element.findall(f'.//{{{ns_w}}}drawing')
        if not text and not drawings:
            insert_idx = j
            break

    insertions.append((insert_idx, png_path, caption, heading_text))
    print(f'  {heading_text}: will insert at P{insert_idx}')

print(f'Found {len(insertions)} insertion positions')

# ============================================================
# Step 3: Do all insertions
# ============================================================
print()
print('='*60)
print('Step 3: Inserting images...')
print('='*60)

inserted = 0
for insert_idx, png_path, caption, heading_text in sorted(insertions, key=lambda x: x[0], reverse=True):
    try:
        p_insert = doc.paragraphs[insert_idx]

        # Add image to paragraph
        run = p_insert.add_run()
        run.add_picture(png_path, width=Inches(5.5))

        # Create caption
        cap_elem = etree.Element(f'{{{ns_w}}}p')
        cap_pPr = etree.SubElement(cap_elem, f'{{{ns_w}}}pPr')
        cap_jc = etree.SubElement(cap_pPr, f'{{{ns_w}}}jc')
        cap_jc.set(f'{{{ns_w}}}val', 'center')
        cap_r = etree.SubElement(cap_elem, f'{{{ns_w}}}r')
        cap_rPr = etree.SubElement(cap_r, f'{{{ns_w}}}rPr')
        cap_sz = etree.SubElement(cap_rPr, f'{{{ns_w}}}sz')
        cap_sz.set(f'{{{ns_w}}}val', '21')
        cap_t = etree.SubElement(cap_r, f'{{{ns_w}}}t')
        cap_t.text = caption
        cap_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        # Insert caption after image paragraph
        p_insert._element.addnext(cap_elem)

        print(f'  {heading_text}: inserted at P{insert_idx}')
        inserted += 1
    except Exception as e:
        print(f'  {heading_text}: FAILED - {e}')

print(f'\nInserted: {inserted}/{len(insertions)}')

# Save
output = '平菇热泵烘干房设计_广州版_CAD附图.docx'
doc.save(output)
print(f'Saved to: {output}')

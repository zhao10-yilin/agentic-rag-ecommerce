#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Create translated Paper 2: 高温预处理联合热风干燥对杏鲍菇蛋白结构与性质的影响
Based on: Effects of high-temperature pretreatment combined with hot-air drying
on the structure and properties of P. eryngii proteins
Published in: International Journal of Biological Macromolecules (2024)
DOI: https://doi.org/10.1016/j.ijbiomac.2024.138306
"""

import docx
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import sys, os, zipfile, shutil

sys.stdout.reconfigure(encoding='utf-8')

# Create new document
doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5

# Configure heading styles
for name, size, bold in [
    ('Heading 1', Pt(16), True),
    ('Heading 2', Pt(12), True),
    ('Heading 3', Pt(12), False),
]:
    h_style = doc.styles[name]
    h_font = h_style.font
    h_font.size = size
    h_font.bold = bold
    if name in ['Heading 2', 'Heading 3']:
        h_font.name = '黑体'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h_style.paragraph_format.line_spacing = 1.5


def add_centered_para(text, font_name='黑体', font_size=Pt(22), bold=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = 0
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return para


def add_body_para(text, font_name='宋体', font_size=Pt(12), bold=False, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return para


def add_heading_para(text, level=1):
    heading = doc.add_heading(text, level=level)
    if level >= 2:
        heading.paragraph_format.line_spacing = 1.5
        for run in heading.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def add_keywords_para(text):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return para


def add_image_from_file(image_path, width_inches=5.5):
    """Add image from file to document"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run()
        try:
            run.add_picture(image_path, width=Inches(width_inches))
        except Exception as e:
            run.text = f'[图片: {os.path.basename(image_path)} - 插入失败: {e}]'
            run.font.size = Pt(9)
        return para
    return None


def add_figure_caption(text):
    """Add figure caption"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para


def add_table_caption(text):
    """Add table caption"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(10.5)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return para


# Set up image paths
IMG_DIR = 'extracted_images_paper2/word/media'

# ============================================================
# TITLE PAGE
# ============================================================
add_centered_para('本科毕业设计（论文）译文', '黑体', Pt(26), False)
doc.add_paragraph()

add_centered_para('高温预处理联合热风干燥对杏鲍菇蛋白\n结构与性质的影响研究', '黑体', Pt(22), True)
doc.add_paragraph()

add_centered_para(
    '原文标题：Effects of high-temperature pretreatment combined with\n'
    'hot-air drying on the structure and properties of P. eryngii proteins',
    '宋体', Pt(14), False)
add_centered_para(
    '原文作者：Deqing Wang, Dianbin Su, Huihui Xu, Xiaofeng Chen, Weiqiao Yang, et al.',
    '宋体', Pt(14), False)
add_centered_para(
    '原文出处：International Journal of Biological Macromolecules (2024)\n'
    'DOI: https://doi.org/10.1016/j.ijbiomac.2024.138306',
    '宋体', Pt(14), False)
doc.add_paragraph()
add_centered_para('2026年 6 月', '黑体', Pt(18), False)

doc.add_page_break()

# ============================================================
# 摘要
# ============================================================
add_heading_para('摘要', 1)

abstract_cn = (
    '杏鲍菇（Pleurotus eryngii）因其丰富的营养价值和独特的质构特性而备受推崇，作为传统蛋白质的替代品'
    '具有广阔的开发前景。本研究系统探讨了不同条件的高温预处理联合热风干燥对杏鲍菇蛋白（PEP）结构'
    '与功能性质的影响。采用四种预处理方式：沸水预处理联合热风干燥（BPHD）、微波预处理联合热风干燥'
    '（MPHD）、红外预处理联合热风干燥（IPHD）和高温热风干燥（HPHD），以真空冷冻干燥（VFD）和新鲜样品'
    '（Fresh-H）作为对照。通过巯基含量、溶解度、表面疏水性、荧光光谱、傅里叶变换红外光谱（FTIR）、'
    'X射线衍射（XRD）、粒径与Zeta电位、微观结构、浊度、持水持油性、乳化性、起泡性、热稳定性、'
    '体外消化模拟及抗氧化活性等多项指标，系统评价了不同预处理方式对杏鲍菇蛋白结构和功能特性的影响。'
    '结果表明，MPHD处理在保持蛋白质良好功能性质方面表现最优，具有较高的乳化性能和稳定性，同时具有'
    '优异的水解度和抗氧化能力。通过对游离氨基酸和挥发性化合物的分析进一步揭示了不同干燥方式对杏鲍菇'
    '风味品质的影响规律。本研究为杏鲍菇蛋白的加工利用提供了理论依据和技术参考。'
)
add_body_para(abstract_cn)
doc.add_paragraph()
add_keywords_para('关键词：杏鲍菇蛋白，高温预处理，热风干燥，蛋白结构，功能性质，体外消化')

doc.add_paragraph()

# ============================================================
# Abstract (English)
# ============================================================
add_heading_para('Abstract', 1)
abstract_en = (
    'P. eryngii (Pleurotus eryngii), commonly known as King Oyster Mushroom, is highly regarded for its rich '
    'nutritional content and unique texture, and shows great potential as a substitute for traditional proteins. '
    'This study systematically investigated the effects of high-temperature pretreatment under different conditions '
    'combined with hot-air drying on the structure and functional properties of P. eryngii proteins (PEP). '
    'Four pretreatment methods were employed: boiling water pretreatment combined with hot air drying (BPHD), '
    'microwave pretreatment combined with hot air drying (MPHD), infrared pretreatment combined with hot air drying '
    '(IPHD), and high-temperature hot air drying (HPHD), with vacuum freeze drying (VFD) and fresh samples '
    '(Fresh-H) as controls. Multiple indicators including sulfhydryl content, solubility, surface hydrophobicity, '
    'fluorescence spectroscopy, FTIR, XRD, particle size and zeta potential, microstructure, turbidity, water- and '
    'oil-holding capacity, emulsifying properties, foaming properties, thermal stability, in vitro digestion '
    'simulation and antioxidant activity were used to systematically evaluate the effects of different pretreatment '
    'methods on the structural and functional properties of PEP. The results showed that MPHD treatment performed '
    'best in maintaining good functional properties of proteins, with superior emulsification performance and '
    'stability, as well as excellent hydrolysis degree and antioxidant capacity. Analysis of free amino acids and '
    'volatile compounds further revealed the influence of different drying methods on the flavor quality of '
    'P. eryngii. This study provides theoretical basis and technical reference for the processing and utilization '
    'of P. eryngii proteins.'
)
add_body_para(abstract_en)
doc.add_paragraph()
add_keywords_para('Key words: P. eryngii proteins, High-temperature pretreatment, Hot air drying, Protein structure, Functional properties, In vitro digestion')

doc.add_page_break()

# ============================================================
# 第一章 绪论
# ============================================================
add_heading_para('第一章 绪论', 1)

add_heading_para('1.1 研究背景', 2)
add_body_para(
    '杏鲍菇（Pleurotus eryngii），俗称King Oyster Mushroom（帝王蚝菇），因其丰富的营养含量和独特的'
    '质构特性而备受推崇[1]。杏鲍菇含有丰富的蛋白质、膳食纤维、多糖以及多种必需氨基酸和矿物质元素，'
    '是一种营养价值极高的食用菌品种。近年来，随着人们对健康和可持续食品的关注日益增加，植物基蛋白质'
    '替代品的需求迅速增长。杏鲍菇蛋白因其良好的氨基酸组成和功能特性，作为传统动物蛋白质的替代品'
    '具有广阔的开发前景。因此，杏鲍菇基蛋白质替代品的研究与开发拓展了植物蛋白的选择范围，满足了消费者'
    '对健康和可持续食品的需求。'
)
add_body_para(
    '蛋白质提取前的预处理对蛋白质的提取效率、品质和功能性质具有重要影响。常见的预处理方法包括'
    '酶水解、化学预处理以及物理预处理（如超声波、微波、热处理等）。其中，热处理因其操作简单、'
    '成本低廉等优势被广泛应用。然而，不同的热处理条件（如温度、时间、加热方式）对蛋白质结构和功能'
    '性质的影响机制存在显著差异。已有研究表明，短期高温处理对山羊乳清蛋白的损伤程度低于长时间低温处理[8]。'
    '干燥方法对蛋白质结构和功能性质也有显著影响，体现在不同干燥方法对蛋白质修饰效果的巨大差异上[9,10]。'
)
add_body_para(
    '在已有研究范围内，热处理方式的选择相对单一，比较多种热处理方式对蛋白质结构和功能性质影响的研究较少。'
    '因此，本研究旨在系统比较不同高温预处理方式（沸水、微波、红外）联合热风干燥对杏鲍菇蛋白结构、'
    '功能性质和体外消化特性的影响，为杏鲍菇蛋白的加工利用提供全面的理论基础和技术指导。'
)

add_heading_para('1.2 研究目的与内容', 2)
add_body_para(
    '本研究的主要目的是：（1）探究不同高温预处理联合热风干燥方式对杏鲍菇蛋白结构特性的影响，'
    '包括巯基和二硫键含量、二级结构、晶体结构、粒径与Zeta电位、微观形貌等指标的变化规律；'
    '（2）系统评价不同预处理方式对杏鲍菇蛋白功能性质的影响，包括溶解度、表面疏水性、持水持油性、'
    '乳化性、起泡性等加工功能特性；'
    '（3）研究不同预处理方式对杏鲍菇蛋白热稳定性和体外消化特性的影响，评估其在食品加工和人体消化'
    '过程中的表现；'
    '（4）分析不同预处理方式对杏鲍菇游离氨基酸组成和挥发性风味化合物的影响，为产品风味品质的调控'
    '提供依据。'
)

doc.add_page_break()

# ============================================================
# 第二章 材料与方法
# ============================================================
add_heading_para('第二章 材料与方法', 1)

add_heading_para('2.1 材料与试剂', 2)
add_body_para(
    '新鲜杏鲍菇（Pleurotus eryngii）购自中国山东省淄博市。实验前去除菌盖和菌柄末端，'
    '切成约20 mm见方的小块。所有化学试剂均为分析纯级别，实验用水为超纯水。'
)

add_heading_para('2.2 设备与干燥处理', 2)

add_heading_para('2.2.1 干燥设备', 3)
add_body_para(
    '杏鲍菇的预处理和热风干燥采用自主研发的ORW4.0S-5000R型微波红外热风滚床干燥器（MIHRBD）进行。'
    '该设备集成了微波、红外和热风三种干燥方式，能够实现精确的温度控制和多种干燥模式的组合操作。'
    '设备的具体结构详见原文图1。'
)

add_heading_para('2.2.2 高温预处理与热风干燥工艺', 3)
add_body_para(
    '基于课题组前期项目的研究成果、预实验结果以及食用菌干制品企业的常规工作条件，确定了以下预处理条件：'
    '（1）沸水预处理联合热风干燥（BPHD）：样品在沸水中处理3分钟后进行热风干燥；'
    '（2）微波预处理联合热风干燥（MPHD）：样品在微波功率300 W下处理2分钟后进行热风干燥；'
    '（3）红外预处理联合热风干燥（IPHD）：样品在红外功率500 W下处理3分钟后进行热风干燥；'
    '（4）高温热风干燥（HPHD）：样品直接在80℃下进行热风干燥，无额外预处理。'
    '以真空冷冻干燥（VFD）样品和新鲜样品（Fresh-H）作为对照。低温热风干燥过程参照之前使用的方法进行'
    '监控[15]。简要而言，每隔10分钟取出样品称重，直至水分含量降至安全贮藏水平（约10%以下）。'
)

add_heading_para('2.3 蛋白质的测定', 2)
add_body_para(
    '杏鲍菇蛋白（PEP）的提取采用Osborne分级提取法，具体提取流程如图1所示。提取得到的蛋白组分'
    '合并后进行冷冻干燥保存，用于后续各项指标的测定。蛋白质含量的测定采用凯氏定氮法。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image5.jpeg'), 5.0)
add_figure_caption('图1 杏鲍菇蛋白提取流程图 (Fig. 1. Extraction Flowchart of P. eryngii proteins)')

add_heading_para('2.3.1 巯基含量测定', 3)
add_body_para(
    '采用Ellman试剂法测定杏鲍菇蛋白的总游离巯基、总巯基及二硫键含量。所用仪器和详细方法见原文附表S2。'
)

add_heading_para('2.3.2 溶解度测定', 3)
add_body_para(
    '杏鲍菇蛋白的溶解度表示为上清液中蛋白质浓度与样品中总蛋白质含量的比值，所用仪器和方法见原文附表S2。'
)

add_heading_para('2.3.3 表面疏水性（H₀）测定', 3)
add_body_para(
    '杏鲍菇蛋白表面疏水性的测定采用ANS荧光探针法，所用仪器和方法见原文附表S2。'
)

add_heading_para('2.3.4 荧光光谱分析', 3)
add_body_para(
    '杏鲍菇蛋白的内源荧光光谱采用荧光分光光度计进行分析，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.5 傅里叶变换红外光谱（FTIR）分析', 3)
add_body_para(
    '杏鲍菇蛋白的二级结构通过FTIR进行分析，仪器和方法见原文附表S2。对酰胺I带（1700-1600 cm⁻¹）'
    '进行傅里叶去卷积处理，计算各二级结构（α-螺旋、β-折叠、β-转角和无规卷曲）的相对含量。'
)

add_heading_para('2.3.6 X射线衍射（XRD）分析', 3)
add_body_para(
    '杏鲍菇蛋白的XRD图谱采用X射线衍射仪进行观察，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.7 粒径与Zeta电位测定', 3)
add_body_para(
    '蛋白质的粒径和Zeta电位采用纳米粒度及Zeta电位分析仪进行分析，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.8 微观结构观察', 3)
add_body_para(
    '采用扫描电子显微镜（SEM）观察杏鲍菇蛋白的微观形貌，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.9 浊度测定', 3)
add_body_para(
    '蛋白质的浊度采用紫外-可见分光光度计进行分析，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.10 持水性（WHC）与持油性（OHC）测定', 3)
add_body_para(
    '杏鲍菇蛋白的持水性和持油性参照文献方法进行测定，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.11 乳化性与起泡性测定', 3)
add_body_para(
    '杏鲍菇蛋白的乳化活性指数（EAI）、乳化稳定性指数（ESI）、起泡能力（FC）和起泡稳定性（FS）'
    '参照文献方法进行测定，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.12 热稳定性分析', 3)
add_body_para(
    '杏鲍菇蛋白的热稳定性采用热重分析仪（TGA）进行测定，仪器和方法见原文附表S2。'
)

add_heading_para('2.3.13 体外消化模拟', 3)
add_body_para(
    '体外消化模拟包括胃消化阶段和肠消化阶段。胃消化阶段在0、30、60、90和120分钟取样，'
    '肠消化阶段在150、180、210、240、270和300分钟取样。分别收集各时间点的酶解液用于后续分析。'
)

add_heading_para('2.3.14 抗氧化活性测定', 3)
add_body_para(
    '酶解液的DPPH和ABTS自由基清除率参照文献方法进行测定，仪器和详细方法见原文附表S2。'
)

add_heading_para('2.4 游离氨基酸（FAA）分析', 2)
add_body_para(
    '干燥杏鲍菇的氨基酸组成采用氨基酸自动分析仪进行测定，仪器和方法见原文附表S2。'
)

add_heading_para('2.5 挥发性化合物分析', 2)
add_body_para(
    '采用顶空-固相微萃取（HS-SPME）结合气相色谱-质谱联用（GC-MS）的方法测定不同干燥处理后'
    '杏鲍菇的挥发性化合物。以1-癸醇为内标物进行定量分析。'
)

add_heading_para('2.6 统计分析', 2)
add_body_para(
    '每个实验重复三次，结果以平均值±标准差表示。采用Origin软件进行图表绘制。采用SPSS软件进行'
    '单因素方差分析（ANOVA），以Duncan多重比较检验进行显著性分析（P < 0.05）。'
)

doc.add_page_break()

# ============================================================
# 第三章 结果与讨论
# ============================================================
add_heading_para('第三章 结果与讨论', 1)

add_heading_para('3.1 巯基分析', 2)
add_body_para(
    '巯基基团以多种方式对蛋白质的功能性质产生重要影响，包括参与二硫键的形成、氧化还原反应以及'
    '修复氧化损伤的蛋白质。二硫键是连接蛋白质中不同肽链的共价键，通过稳定蛋白质的三维空间结构'
    '来维持其原有的功能和性质。如图2A所示，与Fresh-H相比，经不同预处理后杏鲍菇蛋白的游离巯基含量'
    '发生了显著变化。'
)
add_body_para(
    '研究发现，HPHD处理样品具有较低的二硫键含量。有趣的是，虽然MPHD样品中的巯基含量低于HPHD，'
    '但是MPHD的二硫键含量反而更低。这可能是由于微波处理的特殊加热机制导致了蛋白质分子内部'
    '巯基-二硫键交换反应的差异。不同预处理方式对蛋白质中巯基和二硫键含量的影响差异显著，'
    '反映了不同热处理方式对蛋白质结构修饰的不同作用机制。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image6.jpeg'), 5.5)
add_figure_caption('图2 杏鲍菇蛋白的巯基含量(A)、溶解度(B)、表面疏水性(C)和荧光光谱(D) (Fig. 2)')

add_heading_para('3.2 溶解度分析', 2)
add_body_para(
    '热处理对蛋白质溶解度的影响与蛋白质与水分子之间的水合程度以及热处理过程中蛋白质的聚集程度有关[19]。'
    '如图2B所示，不同预处理方式对杏鲍菇蛋白的溶解度产生了显著影响。VFD处理的蛋白保持了较高的溶解度，'
    '这是由于低温真空环境减少了蛋白质的热变性和聚集。相比之下，经高温预处理后的蛋白溶解度均有'
    '不同程度的降低，其中BPHD处理的蛋白溶解度降低最为明显，这可能与沸水处理导致的蛋白质严重变性'
    '和不可逆聚集有关。MPHD处理的蛋白溶解度相对较高，表明微波预处理对蛋白溶解度的不利影响较小，'
    '这可能归因于微波加热的均匀性和较短的处理时间。'
)

add_heading_para('3.3 表面疏水性（H₀）分析', 2)
add_body_para(
    '表面疏水性反映了蛋白质分子表面疏水基团的变化，可以指示蛋白质分子空间构象的改变。'
    '如图2C所示，不同预处理后杏鲍菇蛋白的表面疏水性发生了显著变化。与VFD相比，经热处理后蛋白的'
    '表面疏水性普遍增加，这是由于热处理导致蛋白质分子解析叠，暴露出原本埋藏在分子内部的疏水基团。'
    '其中BPHD处理的蛋白表面疏水性最高，表明沸水处理导致了最显著的蛋白质结构展开。MPHD和IPHD处理的'
    '蛋白表面疏水性增加适中，说明这两种预处理方式对蛋白质结构的影响相对温和。'
)

add_heading_para('3.4 内源荧光光谱分析', 2)
add_body_para(
    '在天然蛋白质中，产生荧光的氨基酸主要是色氨酸、酪氨酸和苯丙氨酸[22]。荧光光谱主要取决于'
    '色氨酸残基对荧光强度的贡献，对色氨酸微环境的极性变化非常敏感。如图2D所示，不同预处理方式'
    '对杏鲍菇蛋白的荧光光谱产生了显著影响。VFD处理的蛋白表现出最高的荧光强度，说明其色氨酸残基'
    '主要埋藏在蛋白质分子内部的疏水环境中。随预处理温度的升高，荧光强度逐渐降低并伴随最大发射波长'
    '的红移，表明蛋白质结构展开，色氨酸残基暴露于更极性的溶剂环境中。BPHD处理表现出最低的荧光强度'
    '和最显著的红移，反映了最剧烈的蛋白质构象变化。'
)

add_heading_para('3.5 结构分析', 2)

add_heading_para('3.5.1 FTIR光谱分析', 3)
add_body_para(
    '杏鲍菇蛋白的FTIR原始光谱如图3A所示。3600-3200 cm⁻¹之间的谱带与蛋白质本身羟基的伸缩振动有关。'
    '3000-2800 cm⁻¹之间的谱带归属于C-H伸缩振动。经预处理后，这些谱带的峰强度有所降低，因为高温'
    '可能导致蛋白质分子裂解和肽键断裂，导致贡献基团数量减少。酰胺II带（C-N伸缩振动）在1550 cm⁻¹处'
    '有吸收峰，经预处理后表现出轻微的红移和峰强度的降低。'
)
add_body_para(
    '对酰胺I带（1700-1600 cm⁻¹）范围进行傅里叶去卷积处理，计算蛋白质的二级结构含量。'
    '表1显示了四种蛋白质二级结构的含量。结果表明，不同预处理方式对杏鲍菇蛋白二级结构的影响'
    '差异显著。VFD处理的蛋白保持了较高的α-螺旋和β-折叠含量，表明其天然结构保持较好。'
    '经热处理后，α-螺旋含量普遍降低，β-折叠和无规卷曲含量增加，反映了蛋白质分子从有序结构向'
    '无序结构的转变。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image7.png'), 2.3)
add_image_from_file(os.path.join(IMG_DIR, 'image8.png'), 2.3)
add_figure_caption('图3 杏鲍菇蛋白的FTIR光谱(A)、反卷积FTIR光谱(B)、XRD(C)、粒径(D)和Zeta电位(E) (Fig. 3)')

add_table_caption('表1 杏鲍菇蛋白二级结构含量 (Table 1)')
add_body_para(
    '（注：表格数据详见原文Table 1，包括α-螺旋、β-折叠、β-转角和无规卷曲四种二级结构的百分比含量）',
    '宋体', Pt(10.5), False, False)

add_heading_para('3.5.2 XRD分析', 3)
add_body_para(
    '通过XRD测定物质的晶体结构，包括精确鉴定晶相组分和研究加工条件下的结构变化，可以帮助理解'
    '蛋白质在加工过程中的结构演变。不同预处理方式对蛋白的XRD图谱产生了显著影响。VFD处理蛋白的'
    'XRD图谱显示出较宽的弥散峰，表明其以无定形结构为主。经热预处理后，某些衍射峰的强度发生变化，'
    '反映了蛋白质结晶度和分子有序性的改变。已有研究在微波和冷冻干燥技术中也证明了不同干燥和预处理方法'
    '对各种物质结晶度的类似影响[29,30]。这些研究表明，涉及热效应的加工过程通常会导致分子有序性的'
    '重新排列。'
)

add_heading_para('3.5.3 粒径与Zeta电位分析', 3)
add_body_para(
    '粒径可以有效反映蛋白质聚集和交联的程度。如图3D所示，与Fresh-H相比，经预处理后杏鲍菇蛋白的'
    '粒径有不同程度的增加。BPHD处理的蛋白粒径增加最为显著，这是由于沸水处理导致了严重的蛋白质变性'
    '和聚集，形成了较大的不溶性聚集体。MPHD和IPHD处理的蛋白粒径增加较小，表明这两种预处理方式'
    '引起的蛋白质聚集程度较轻。'
)
add_body_para(
    'Zeta电位值与胶体分散体系的稳定性相关，反映了邻近带有相同电荷的粒子之间的排斥程度[31]。'
    'Zeta电位的绝对值越大，体系越稳定。如图3E所示，不同预处理方式对蛋白的Zeta电位产生了显著影响。'
    'VFD处理的蛋白具有较高的Zeta电位绝对值，表明其分散体系较为稳定。经热预处理后，Zeta电位的绝对值'
    '普遍降低，表明蛋白质的分散稳定性有所下降。'
)

add_heading_para('3.5.4 微观结构分析', 3)
add_body_para(
    '蛋白质的形态和结构对其功能性质（包括持水持油性、乳化性和起泡性）具有重要影响，直接关系到'
    '蛋白质在食品加工中的应用性能。如图4A的SEM图像所示，不同预处理后杏鲍菇蛋白的微观形貌差异显著。'
    'VFD处理的蛋白呈现出疏松多孔的片层结构，这种结构有利于水的结合和保持。经高温预处理后，'
    '蛋白的微观结构变得更加致密和不规则。BPHD处理的蛋白呈现出高度致密和聚集的块状结构，'
    '与其低溶解度和功能性质的劣化相对应。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image15.jpeg'), 5.5)
add_figure_caption('图4 预处理后杏鲍菇蛋白的SEM图像(A)及预处理对蛋白作用的机理示意图(B) (Fig. 4)')

add_heading_para('3.6 理化指标分析', 2)

add_heading_para('3.6.1 浊度分析', 3)
add_body_para(
    '浊度可以表征蛋白质溶液的光学性质，直观反映蛋白质颗粒在分散体系中的分散状态，指示蛋白质的'
    '聚集状态。如图5A所示，不同预处理方式对蛋白溶液浊度的影响与其粒径变化趋势一致。BPHD处理的'
    '蛋白浊度最高，表明形成了大量的不溶性聚集体。MPHD和IPHD处理的蛋白浊度增加相对较小，与其较好'
    '的分散性相符。'
)

add_heading_para('3.6.2 持水性（WHC）与持油性（OHC）分析', 3)
add_body_para(
    '如图5B所示，与Fresh-H蛋白相比，VFD处理后持水性的增加不显著。BPHD处理后持水性增加最为显著，'
    '其次是MPHD、IPHD和HPHD。这是因为蛋白质表面的带电和极性基团对水具有强亲和力，热处理导致蛋白'
    '结构展开，使更多的极性基团暴露于分子表面，从而增强了与水的结合能力。然而，过度的热处理（BPHD）'
    '导致蛋白质严重聚集，部分抵消了结构展开对持水性的积极作用。在持油性方面，各处理组之间的变化趋势'
    '与持水性有所不同，反映了蛋白质表面疏水基团和亲水基团暴露比例的变化对水和油的结合产生了不同影响。'
)

add_heading_para('3.6.3 乳化性与起泡性分析', 3)
add_body_para(
    'EAI表示蛋白质形成乳液的能力，而ESI反映乳液在给定时间范围内保持其分散状态和抵抗分离的能力。'
    '如图5C所示，与Fresh-H相比，不同预处理后蛋白的EAI和ESI发生了显著变化。MPHD处理的蛋白表现出'
    '最高的EAI，表明微波预处理有利于蛋白质在油-水界面的吸附和扩散。BPHD处理的蛋白EAI最低，'
    '这是由于蛋白质的热聚集和干燥过程中糖与蛋白质反应的加深[9]。有趣的是，BPHD的EAI较低，'
    '但ESI并不低，这可能与其形成的高度交联的界面蛋白膜有关。'
)
add_body_para(
    '泡沫的形成取决于蛋白质形成界面膜、维持悬浮气泡和减缓聚并速率的能力。如图5D所示，与Fresh-H相比，'
    '不同预处理后蛋白的FC和FS均有变化。MPHD和IPHD处理的蛋白FC较高，这可能与其较好的溶解度和'
    '适度的结构展开有关。FS表现出与FC不同的趋势，IPHD处理的蛋白具有良好的起泡能力但FS较差，'
    '表明其形成的泡沫虽然数量多但稳定性不足。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image16.png'), 1.8)
add_image_from_file(os.path.join(IMG_DIR, 'image21.png'), 1.8)
add_figure_caption('图5 杏鲍菇蛋白的浊度(A)、持水持油性(B)、乳化性(C)和起泡性(D) (Fig. 5)')

add_heading_para('3.7 热稳定性分析', 2)
add_body_para(
    '对预处理后的杏鲍菇蛋白进行热重分析（TGA），主要目的是研究其热稳定性和不同处理方式之间的'
    '潜在结构差异。如图6A所示，所有蛋白样品的热重曲线显示出相似的三阶段热降解模式：第一阶段'
    '（50-150℃）主要为自由水和结合水的蒸发；第二阶段（200-400℃）为蛋白质主链降解和肽键断裂；'
    '第三阶段（400℃以上）为碳化和残余物形成。'
)
add_body_para(
    '值得注意的是，MPHD和IPHD处理的蛋白最终残余率分别为31.19%和29.1%，而BPHD处理蛋白的最终残余率'
    '仅为2.81%。MPHD和IPHD相对较高的残余率表明这些预处理促进了蛋白分子间更强的交联和相互作用，'
    '从而使其更能抵抗热分解[39]。已有研究也提供了证据，支持此类预处理在改善蛋白质热稳定性方面的作用。'
)

add_heading_para('3.8 体外消化分析', 2)

add_heading_para('3.8.1 蛋白质水解度分析', 3)
add_body_para(
    '如图6B所示，不同方法处理的杏鲍菇蛋白的水解趋势相似。总体而言，在胃消化阶段水解度缓慢增加，'
    '进入肠消化阶段后水解迅速。在肠消化阶段，MPHD处理的蛋白具有最高的水解度。与VFD相比，'
    '其他预处理方法中蛋白的消化率均有提高。消化240分钟后，不同处理组的蛋白消化率趋于稳定。'
    'MPHD处理的蛋白优异的水解度可能与其适度的结构展开和较低的聚集程度有关，这有利于消化酶'
    '对蛋白质肽键的可及性和酶解作用。'
)

add_heading_para('3.8.2 抗氧化活性分析', 3)
add_body_para(
    '抗氧化能力的评价通过测定样品的自由基清除能力来实现，用于比较不同样品、研究加工对其的影响，'
    '并对功能性食品的开发具有指导意义。如图6C和D所示，不同预处理方法获得的杏鲍菇蛋白消化产物的'
    '抗氧化活性表现出明显差异。在胃消化阶段，DPPH清除率随消化时间的延长而增加。进入肠消化阶段后，'
    'DPPH清除率逐渐趋于稳定，表明体系中抗氧化物质的生成与消耗之间达到了动态平衡。'
)
add_body_para(
    '比较不同杏鲍菇蛋白的抗氧化活性发现，预处理后的蛋白抗氧化活性相比VFD有所增强。这可能归因于'
    '热处理导致蛋白质结构的部分展开，释放出具有抗氧化活性的肽段和氨基酸残基。MPHD处理的蛋白在肠消化'
    '阶段表现出最高的DPPH和ABTS自由基清除率，与其最高的水解度一致，表明适度预处理有利于释放更多'
    '的抗氧化肽段。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image22.png'), 5.5)
add_figure_caption('图6 杏鲍菇蛋白的TGA分析(A)、水解度(B)、DPPH(C)和ABTS(D)自由基清除能力 (Fig. 6)')

add_heading_para('3.9 游离氨基酸（FAA）分析', 2)
add_body_para(
    '如表2所示，不同处理方法对杏鲍菇中总游离氨基酸和各单个游离氨基酸含量产生了显著影响。'
    '与新鲜杏鲍菇相比，经不同干燥处理后总游离氨基酸含量均有所下降。BPHD、HPHD和IPHD处理样品'
    '的总游离氨基酸含量分别降低了80.44%、60.22%、62.2%和67.66%，必需氨基酸降低了64.62%至84.06%。'
    '游离氨基酸的损失主要是由于干燥过程中游离氨基酸与可溶性糖之间发生的美拉德反应所致。'
)
add_body_para(
    '与新鲜杏鲍菇相比，VFD处理中游离氨基酸和必需氨基酸的含量保持相对稳定。这是因为冷冻干燥的'
    '低温真空环境减少了美拉德反应和其他氨基酸降解反应的发生。在热处理组中，MPHD处理的游离氨基酸'
    '保留率相对较高，表明微波预处理的短时加热有利于减少氨基酸的损失。'
)

add_table_caption('表2 杏鲍菇游离氨基酸含量 (Table 2)')
add_body_para(
    '（注：表格详细数据（包括天冬氨酸、谷氨酸、丝氨酸等17种游离氨基酸的含量）详见原文Table 2）',
    '宋体', Pt(10.5), False, False)

add_heading_para('3.10 挥发性化合物分析', 2)
add_body_para(
    '挥发性化合物通过多种途径形成，包括美拉德反应、长链化合物的降解、不饱和脂肪酸的氧化以及与蛋白质、'
    '肽和氨基酸的相互作用。醇类在杏鲍菇的挥发性化合物中含量最高，其前体主要是多不饱和脂肪酸。'
    '1-辛烯-3-醇是主要的挥发性化合物，其次是3-辛醇和1-辛醇等。在干燥过程中产生了新的醇类化合物，'
    '增强了产品的风味。'
)
add_body_para(
    '酯类是食品中甜味和微油香气的首要来源。干燥后，酯类化合物的含量显著增加，其中HPHD处理的'
    '酯类含量最高。MPHD和IPHD处理的酯类含量也有明显提高，但BPHD处理的酯类增加较少，'
    '可能是由于过度的热处理导致酯类化合物的挥发损失。醛类在食用菌中含量丰富且气味阈值低，'
    '对整体挥发性化合物贡献显著。碳链为C5至C9的醛类主要来源于脂肪酸的氧化降解。'
)
add_body_para(
    '共检测到11种碳氢化合物。HPHD处理中烷烃类含量最多，IPHD处理中含量也相对较高。'
    '不同干燥方法产生了新的烷烃类化合物。虽然碳氢化合物通常具有较高的气味阈值，但由于其含量较大，'
    '对产品的整体风味仍有不可忽视的贡献。'
)

add_table_caption('表3 杏鲍菇挥发性化合物含量 (Table 3)')
add_body_para(
    '（注：表格详细数据（包括醇类、酯类、醛类、酮类、碳氢化合物等各类挥发性化合物的定性和定量结果）'
    '详见原文Table 3）',
    '宋体', Pt(10.5), False, False)

add_heading_para('3.11 风味PCA分析', 2)
add_body_para(
    '游离氨基酸和挥发性化合物的主成分分析（PCA）结果如图7所示。PCA1和PCA2的贡献率分别为66.24%和'
    '18.97%，累计贡献率达到85.21%，说明前两个主成分可以解释绝大部分数据变异。不同处理组的蛋白'
    '在PCA得分图中的分布可以直观反映其风味特征的相似性和差异性。'
)
add_body_para(
    'MPHD、IPHD和HPHD处理的样品在PCA得分图中分布较为接近，表明这三种处理方式对杏鲍菇风味特征'
    '的影响较为相似。Fresh和BPHD样品与其他样品在PCA1和PCA2上均表现出较大差异，这反映了新鲜样品'
    '和剧烈热处理样品与其他处理组在氨基酸和挥发性化合物轮廓上的显著区别。PCA分析结果为不同干燥方式'
    '对杏鲍菇风味品质影响的综合评价提供了直观的可视化依据。'
)

add_image_from_file(os.path.join(IMG_DIR, 'image23.png'), 5.0)
add_figure_caption('图7 杏鲍菇风味的PCA分析 (Fig. 7. PCA analysis of P. eryngii flavor)')

doc.add_page_break()

# ============================================================
# 第四章 结论
# ============================================================
add_heading_para('第四章 结论', 1)

add_body_para(
    '本研究系统探讨了不同条件的高温预处理联合热风干燥对杏鲍菇蛋白（PEP）结构和功能性质的影响，'
    '并对游离氨基酸和挥发性化合物进行了分析，得出以下主要结论：'
)
add_body_para(
    '（1）不同预处理方式对杏鲍菇蛋白的结构特性产生了显著影响。巯基-二硫键含量分析表明，'
    '热处理导致蛋白质分子内和分子间二硫键的重排；FTIR和XRD分析揭示了蛋白二级结构和晶体结构'
    '的显著变化；粒径和微观结构分析证实了热处理导致蛋白聚集和微观形貌的改变。其中，微波预处理'
    '（MPHD）对蛋白结构的影响最为温和，能够较好地保持蛋白的天然构象。'
)
add_body_para(
    '（2）功能性质方面，MPHD处理的蛋白在溶解性、乳化性和起泡性等方面表现最优，具有较高的乳化性能'
    '和稳定性。BPHD处理因导致蛋白严重变性和聚集，其功能性质劣化最为明显。不同预处理方式对蛋白'
    '功能性质的影响与其对蛋白结构修饰的程度密切相关。'
)
add_body_para(
    '（3）体外消化和抗氧化分析表明，MPHD处理的蛋白具有最优异的水解度和抗氧化能力，使其成为'
    '杏鲍菇蛋白加工的一种新型优选方案。游离氨基酸和挥发性化合物的分析证实了不同干燥方式对'
    '杏鲍菇风味品质具有显著影响，PCA分析直观展示了各处理组风味特征的相似性和差异性。'
)
add_body_para(
    '（4）综合分析认为，微波预处理联合热风干燥（MPHD）是一种较优的杏鲍菇干燥前处理方式，能够在'
    '保证干燥效率的同时，较好地保持蛋白的功能性质和营养品质，为杏鲍菇的工业化干燥加工提供了重要的'
    '理论依据和实践指导。'
)

doc.add_page_break()

# ============================================================
# 参考文献
# ============================================================
add_heading_para('参考文献', 1)

references = [
    '[1] Wang D, Su D, Xu H, Chen X, Yang W, et al. Effects of high-temperature pretreatment combined with hot-air drying on the structure and properties of P. eryngii proteins. Int J Biol Macromol 2024. https://doi.org/10.1016/j.ijbiomac.2024.138306.',
    '[2] Liu Z-W, Tang P-P, Zhang Y-X, Cheng J-H, Aadil RM, Liu X-B. Preventing thermal aggregation of ovalbumin through dielectric-barrier discharge plasma treatment and enhancing its emulsification properties. Int J Biol Macromol 2024;267:131578.',
    '[3] Iida M, Tsuda S, Kikuchi M, Samoto M, Adachi N, Nakamura A. Extraction of water-soluble polysaccharides from lupin beans and their function of protein dispersion and stabilization under acidic conditions. Int J Biol Macromol 2024;278:134664.',
    '[4] Ganjeh AM, Pinto CA, Casal S, Saraiva JA. The effects of pressure-based processing technologies on protein oxidation. Food Biosci 2024;59:103963.',
    '[5] Koop J, Merz J, Schembecker G. Hydrophobicity, amphilicity, and flexibility: relation between molecular protein properties and the macroscopic effects of surface activity. J Biotechnol 2021;334:50-60.',
    '[6] Joshi S, Asthana S, Pandey MD, Pandey R. Thiazole-based metalloligands and their heteroleptic MOFs: chromogenic and fluorometric detection of tryptophan and Z-l-phenylalanine. Int J Biol Macromol 2024.',
    '[7] Han C, Zheng Y, Huang S, Xu L, Zhou C, Sun Y, et al. Exploring the binding mechanisms of thermally and ultrasonically induced molten globule-like β-lactoglobulin: spectroscopic techniques and molecular simulation. Int J Biol Macromol 2024;263:130300.',
    '[8] Chakraborty S, Paidi MK, Dhinakarasamy I, Sivakumar M, Clements C, Thirumurugan NK, et al. Adaptive mechanism of the marine bacterium Pseudomonas sihuiensis-BFB-6S towards pCO₂ variation. Int J Biol Macromol 2024;261:129860.',
    '[9] Jiang B, Yue H, Fu X, Wang J, Feng Y, Li D, et al. One-step high efficiency separation of prolyl endopeptidase from Aspergillus niger and its application. Int J Biol Macromol 2024;271:132582.',
    '[10] Sousa NFC, Santos MPF, Barbosa RP, Bonomo RCF, Veloso CM, Souza Júnior EC. Pepsin immobilization on activated carbon and functionalized with glutaraldehyde and genipin for the synthesis of antioxidant peptides of goat casein. Food Res Int 2024;186:114161.',
    '[11] Zhang R, Zhang J, Zou B, Ren C, Na X, Xu X, et al. Mild alkalinity preheating treatment regulates the heat and ionic strength co-tolerance of whey protein aggregates. Food Res Int 2024;193:114845.',
    '[12] Ma C, Wan Q, Song J, Hao T, Xia S, Shen S, et al. Ultrasound-assisted pH shift-induced interfacial remodeling for enhancing soluble yeast protein content. Food Hydrocoll 2024;149:109521.',
    '[13] Mahalaxmi S, Himashree P, Malini B, Sunil CK. Effect of microwave treatment on the structural and functional properties of proteins in lentil flour. J Agric Food Res 2022;1:100147.',
    '[14] Yin Y, Liu W, Li L, Cao W, Chen J, Zhao L, et al. Microwave freeze-drying characteristics and crosslinking behavior of wheat starch-laurel acid complex. Int J Biol Macromol 2024.',
    '[15] Hussain Badar I, Wang Z, Chen Q, Liu Q, Ma J, Liu H, et al. Ultrasonic enhancement of structural and emulsifying properties of heat-treated soy protein isolate nanoparticles. Food Chem 2024;442:138469.',
    '[16] Maryniak NZ, Sancho AI, Nielsen SD, Larsen LB, Gao Y, Bggh KL, et al. Enzymatic hydrolysis and extensive heat treatment induce distinct modifications of cow\'s and camel milk proteins. LWT 2024;191:115591.',
    '[17] Anema SG. The turbidity of heated milk in relation to particle size and protein distributions. Int Dairy J 2023;147:105771.',
    '[18] Cao Y, Sun M, Huang T, Zhu Z, Huang M. Effects of heat sterilization on protein physicochemical properties and release of metabolites of braised chicken after in vitro digestion. Food Chem 2024.',
    '[19] Dong X, Woo MW, Quek SY. The physicochemical properties, functionality, and digestibility of hempseed protein isolate as impacted by spray drying and freeze drying. Food Chem 2024;433:137310.',
    '[20] Lee S, Han S, Jo K, Jung S. The impacts of freeze-drying-induced stresses on the quality of meat and aquatic products. Food Chem 2024;459:140437.',
    '[21] Zhang G, Li Y, Song T, Bao M, Li Y, Li X. Colloids Surf B Biointerfaces 2019;181:688-695.',
    '[22] Zhu R, Jones OG. Effect of high acyl gellan gum and pH on the structural and foaming properties of heated whey protein suspensions. Food Chem 2024;449:139255.',
    '[23] Han W, Shi W, Gong D, Zhang G. Improvement of solubility, emulsification property and stability of potato protein by pH-shifting combined with microwave treatment. Food Biosci 2023;56:103301.',
    '[24] Singh S, Bhat HF, Kumar S, Lone AB, Aadil RM, Aït-Kaddour A, et al. Ultrasonication and microwave pre-treated locust protein hydrolysates enhanced the storage stability of meat. Food Chem 2024.',
    '[25] Lotfy SN, Fadel HHM, El-Ghorab AH, Shaheen MS. Stability of encapsulated beef-like flavourings prepared from enzymatically hydrolysed mushroom proteins. Food Chem 2024.',
    '[26] Xie H, Zhang L, Chen Q, Hu J, Zhang P, Xiong H, et al. Combined effects of drying methods and limited enzymatic hydrolysis on the physicochemical and antioxidant properties of rice protein hydrolysates. Food Biosci 2023;52:102427.',
    '[27] Wang Z-Y, Luo Y, Zheng R, Lv S, Li D-Y, Liu Y-X, et al. Effects of air frying on protein digestive properties in scallop adductor muscles. Food Biosci 2024;60:104429.',
    '[28] Na\'thia-Neves G, Alonso E. Valorization of sunflower by-product using microwave-assisted extraction. Food Bioprod Process 2021;125:57-67.',
    '[29] Javed MR, Ahmad Z, Waseem M, Mehmood T, Hussain A, Adil M, et al. Effect of microwave heat processing on nutritional, antioxidant, antinutrient, and sensory indices of soy flour enriched functional noodles. J Agric Food Res 2024;18:101426.',
    '[30] Cao H, Huang Q, Shi J, Guan X, Song H, Zhang Y, et al. Effect of conventional and microwave heating treatment on antioxidant activity of quinoa protein after simulated gastrointestinal digestion. Food Chem 2023;415:135763.',
    '[31] Li Y, He Y, Zhou H, Kinyoro IS, Ma S, Wang J, et al. Ultrasonic-microwave assisted extraction of selenium-rich rice bran protein: a study on its structural and physicochemical properties. LWT 2024;212:116940.',
]

for ref in references:
    add_body_para(ref, 'Times New Roman', Pt(10.5), False, False)

# Save document
output_path = '高温预处理对杏鲍菇蛋白的影响_译文.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Images inserted from: {IMG_DIR}")
